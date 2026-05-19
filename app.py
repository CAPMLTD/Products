import streamlit as st
import anthropic
import os
import json
from datetime import datetime
from io import BytesIO

# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="EU AI Act Model Card Generator",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── Styling ────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
  @import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@300;400;500;600;700&family=DM+Mono:wght@400;500&display=swap');

  html, body, [class*="css"] { font-family: 'DM Sans', sans-serif; }

  .main { background: #f8f9fb; }

  .hero {
    background: linear-gradient(135deg, #0f172a 0%, #1e3a5f 60%, #0f4c81 100%);
    border-radius: 16px;
    padding: 2.5rem 3rem;
    margin-bottom: 2rem;
    color: white;
  }
  .hero h1 { font-size: 2rem; font-weight: 700; margin: 0 0 0.5rem 0; letter-spacing: -0.5px; }
  .hero p  { font-size: 1rem; opacity: 0.75; margin: 0; }

  .risk-card {
    border: 2px solid #e2e8f0;
    border-radius: 12px;
    padding: 1.25rem 1.5rem;
    cursor: pointer;
    transition: all 0.2s;
    background: white;
    margin-bottom: 0.75rem;
  }
  .risk-card:hover { border-color: #3b82f6; background: #eff6ff; }
  .risk-unacceptable { border-left: 5px solid #dc2626; }
  .risk-high         { border-left: 5px solid #f97316; }
  .risk-limited      { border-left: 5px solid #eab308; }
  .risk-minimal      { border-left: 5px solid #22c55e; }

  .section-header {
    background: #0f172a;
    color: white;
    padding: 0.75rem 1.25rem;
    border-radius: 8px 8px 0 0;
    font-weight: 600;
    font-size: 0.95rem;
    letter-spacing: 0.3px;
    margin-top: 1.5rem;
  }
  .section-body {
    border: 1px solid #e2e8f0;
    border-top: none;
    border-radius: 0 0 8px 8px;
    padding: 1.5rem;
    background: white;
    margin-bottom: 0.5rem;
  }

  .compliance-badge {
    display: inline-block;
    padding: 0.25rem 0.75rem;
    border-radius: 999px;
    font-size: 0.75rem;
    font-weight: 600;
    margin-left: 0.5rem;
  }
  .badge-required  { background: #fee2e2; color: #991b1b; }
  .badge-high      { background: #ffedd5; color: #9a3412; }
  .badge-optional  { background: #f0fdf4; color: #166534; }

  .score-bar-outer {
    background: #e2e8f0;
    border-radius: 999px;
    height: 10px;
    margin: 0.5rem 0;
  }
  .score-bar-inner {
    height: 10px;
    border-radius: 999px;
    background: linear-gradient(90deg, #3b82f6, #0f4c81);
  }

  .stTextInput > div > div > input,
  .stTextArea > div > div > textarea,
  .stSelectbox > div > div {
    border-radius: 8px !important;
    border: 1.5px solid #e2e8f0 !important;
    font-family: 'DM Sans', sans-serif !important;
  }
  .stTextInput > div > div > input:focus,
  .stTextArea > div > div > textarea:focus {
    border-color: #3b82f6 !important;
    box-shadow: 0 0 0 3px rgba(59,130,246,0.15) !important;
  }

  .stButton > button {
    border-radius: 8px !important;
    font-family: 'DM Sans', sans-serif !important;
    font-weight: 600 !important;
  }

  .generated-card {
    background: white;
    border: 1px solid #e2e8f0;
    border-radius: 12px;
    padding: 2rem;
    margin-top: 1.5rem;
  }

  .article-tag {
    font-size: 0.7rem;
    font-weight: 600;
    color: #6366f1;
    text-transform: uppercase;
    letter-spacing: 0.5px;
    font-family: 'DM Mono', monospace;
  }

  hr.divider {
    border: none;
    border-top: 1px solid #e2e8f0;
    margin: 1.5rem 0;
  }
</style>
""", unsafe_allow_html=True)

# ── Helpers ────────────────────────────────────────────────────────────────────
def section_header(title, article_ref=None, badge=None):
    badge_html = ""
    if badge == "required":
        badge_html = '<span class="compliance-badge badge-required">Required – High Risk</span>'
    elif badge == "high":
        badge_html = '<span class="compliance-badge badge-high">All Risk Tiers</span>'
    elif badge == "optional":
        badge_html = '<span class="compliance-badge badge-optional">Optional</span>'
    article_html = f'<span class="article-tag" style="float:right;margin-top:2px;">{article_ref}</span>' if article_ref else ""
    st.markdown(f'<div class="section-header">{title}{badge_html}{article_html}</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-body">', unsafe_allow_html=True)

def section_end():
    st.markdown('</div>', unsafe_allow_html=True)

def risk_colour(tier):
    return {"Unacceptable": "#dc2626", "High": "#f97316", "Limited": "#eab308", "Minimal Risk": "#22c55e"}.get(tier, "#64748b")

def compute_score(data, tier):
    """Return a 0-100 completeness score weighted by EU AI Act criticality."""
    required_high = [
        "model_name", "model_type", "developers", "description",
        "data_sources", "data_size", "preprocessing",
        "eval_datasets", "metrics_used", "results_summary",
        "technical_limitations", "fairness_concerns",
        "primary_use_cases", "human_oversight",
        "risk_tier", "risk_justification",
        "data_governance", "accuracy_robustness",
        "post_market_monitoring", "incident_reporting",
        "conformity_assessment_type",
        "out_of_scope_uses",
    ]
    optional = [
        "version", "release_date", "license", "contact",
        "date_range", "known_biases", "demographic_notes",
        "disaggregated_eval", "benchmark_comparisons",
        "failure_modes", "conditions_degrade",
        "recommended_deployment", "intended_users",
        "regulatory_notes", "prohibited_uses", "misuse_scenarios",
        "additional_notes", "cybersecurity_measures",
        "notified_body", "declaration_of_conformity",
    ]
    filled_req = sum(1 for k in required_high if data.get(k, "").strip())
    filled_opt = sum(1 for k in optional if data.get(k, "").strip())
    score = (filled_req / len(required_high)) * 80 + (filled_opt / len(optional)) * 20
    return round(score)

# ── Init session state ─────────────────────────────────────────────────────────
if "step" not in st.session_state:
    st.session_state.step = "wizard"   # wizard → form → result
if "data" not in st.session_state:
    st.session_state.data = {}
if "generated" not in st.session_state:
    st.session_state.generated = ""

# ── Hero ───────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="hero">
  <h1>🛡️ EU AI Act Model Card Generator</h1>
  <p>Generate comprehensive, compliant model cards aligned with EU AI Act Articles 9, 10, 11, 13, 14, 15 & 62 — powered by Claude</p>
</div>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# STEP 1 — RISK CLASSIFICATION WIZARD
# ══════════════════════════════════════════════════════════════════════════════
if st.session_state.step == "wizard":
    st.subheader("Step 1 of 2 — Risk Classification")
    st.markdown("Answer these questions to determine the correct EU AI Act risk tier. This controls which fields are required in your model card.")
    st.markdown("---")

    col1, col2 = st.columns([1.2, 1])

    with col1:
        q1 = st.selectbox(
            "Does the system fall into a prohibited use category (e.g. social scoring by government, real-time remote biometric ID in public spaces, subliminal manipulation)?",
            ["No", "Yes", "Unsure — need guidance"]
        )
        q2 = st.selectbox(
            "Is the system used in a high-risk domain listed in Annex III? (e.g. recruitment, credit scoring, law enforcement, critical infrastructure, education, healthcare, migration, justice)",
            ["No", "Yes", "Unsure — need guidance"]
        )
        q3 = st.selectbox(
            "Does the system interact directly with people in a way that could significantly affect their rights or safety?",
            ["No", "Yes"]
        )
        q4 = st.selectbox(
            "Is the system a general-purpose AI (GPAI) model?",
            ["No", "Yes — standard GPAI", "Yes — systemic risk GPAI (>10²⁵ FLOPs training compute)"]
        )

    with col2:
        st.markdown("#### EU AI Act Risk Tiers")
        st.markdown("""
<div class="risk-card risk-unacceptable">
  <strong style="color:#dc2626">🔴 Unacceptable Risk</strong><br>
  <small>Banned outright. Social scoring, real-time biometric surveillance, subliminal manipulation.</small>
</div>
<div class="risk-card risk-high">
  <strong style="color:#f97316">🟠 High Risk</strong><br>
  <small>Permitted with strict obligations. Annex III domains: healthcare, recruitment, credit, law enforcement, education, migration, justice.</small>
</div>
<div class="risk-card risk-limited">
  <strong style="color:#eab308">🟡 Limited Risk</strong><br>
  <small>Transparency obligations only. Chatbots, deepfakes, emotion recognition systems.</small>
</div>
<div class="risk-card risk-minimal">
  <strong style="color:#22c55e">🟢 Minimal Risk</strong><br>
  <small>No specific obligations beyond voluntary codes of practice. Most AI systems.</small>
</div>
""", unsafe_allow_html=True)

    st.markdown("---")

    # Derive tier
    if q1 in ["Yes", "Unsure — need guidance"]:
        derived_tier = "Unacceptable"
        tier_note = "⚠️ This system may be prohibited under EU AI Act Article 5. A model card alone is insufficient — legal review is required before deployment."
    elif q2 == "Yes":
        derived_tier = "High"
        tier_note = "This system is High Risk under Annex III. All Article 9, 10, 11, 13, 14, 15 obligations apply."
    elif q2 == "Unsure — need guidance":
        derived_tier = "High"
        tier_note = "⚠️ Defaulting to High Risk given uncertainty. Review Annex III with legal counsel."
    elif q3 == "Yes" or "systemic" in q4:
        derived_tier = "Limited"
        tier_note = "Transparency obligations apply under Article 13. Disclose AI nature to users."
    elif q4 != "No":
        derived_tier = "Limited"
        tier_note = "GPAI models have transparency and copyright obligations under Articles 53–55."
    else:
        derived_tier = "Minimal Risk"
        tier_note = "Minimal obligations. Voluntary adherence to codes of practice recommended."

    colour = risk_colour(derived_tier)
    st.markdown(f"""
    <div style="background:{colour}15;border:2px solid {colour};border-radius:10px;padding:1rem 1.5rem;margin-bottom:1rem;">
      <strong style="color:{colour};font-size:1.1rem;">Derived tier: {derived_tier}</strong><br>
      <span style="font-size:0.9rem;color:#374151;">{tier_note}</span>
    </div>
    """, unsafe_allow_html=True)

    override = st.selectbox(
        "Override tier (optional — use if you disagree with the derived classification):",
        ["Use derived tier", "Unacceptable", "High", "Limited", "Minimal Risk"]
    )
    final_tier = derived_tier if override == "Use derived tier" else override

    if st.button("Continue to Model Card Form →", type="primary", use_container_width=True):
        st.session_state.data["risk_tier"] = final_tier
        st.session_state.data["risk_justification"] = tier_note
        st.session_state.step = "form"
        st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# STEP 2 — FULL FORM
# ══════════════════════════════════════════════════════════════════════════════
elif st.session_state.step == "form":
    tier = st.session_state.data.get("risk_tier", "High")
    colour = risk_colour(tier)
    is_high = tier in ["High", "Unacceptable"]

    st.markdown(f"""
    <div style="background:{colour}15;border:2px solid {colour};border-radius:10px;padding:0.75rem 1.25rem;margin-bottom:1.5rem;">
      <strong style="color:{colour};">Risk Tier: {tier}</strong> &nbsp;·&nbsp;
      <span style="font-size:0.85rem;color:#374151;">{st.session_state.data.get('risk_justification','')}</span>
    </div>
    """, unsafe_allow_html=True)

    d = st.session_state.data

    # ── 1 · Model Description ──────────────────────────────────────────────────
    section_header("1 · Model Description", "Art. 11 / Annex IV", "high")
    col1, col2 = st.columns(2)
    with col1:
        d["model_name"]    = st.text_input("Model name *", value=d.get("model_name",""), placeholder="e.g. SeeChange-FFC-Shrink-v2")
        d["model_type"]    = st.text_input("Model type / architecture *", value=d.get("model_type",""), placeholder="e.g. YOLO11 + ResNet classification head")
        d["developers"]    = st.text_input("Developers / Organisation *", value=d.get("developers",""), placeholder="e.g. CA Project Management Services Ltd")
    with col2:
        d["version"]       = st.text_input("Version", value=d.get("version",""), placeholder="e.g. 2.1.0")
        d["release_date"]  = st.text_input("Release date", value=d.get("release_date",""), placeholder="e.g. 2026-05-19")
        d["license"]       = st.text_input("Licence", value=d.get("license",""), placeholder="e.g. Proprietary / Apache 2.0")
        d["contact"]       = st.text_input("Contact email / URL", value=d.get("contact",""), placeholder="e.g. models@ca-pm.co.uk")
    d["description"] = st.text_area("Detailed model description *", value=d.get("description",""), height=120,
        placeholder="Purpose, key capabilities, what problem it solves, notable architectural decisions.")
    section_end()

    # ── 2 · Training Data Characteristics (Art. 10) ───────────────────────────
    section_header("2 · Training Data Characteristics", "Art. 10", "required" if is_high else "high")
    col1, col2 = st.columns(2)
    with col1:
        d["data_sources"]  = st.text_area("Data sources *", value=d.get("data_sources",""), height=100,
            placeholder="List datasets, databases, proprietary sources, annotation providers.")
        d["data_size"]     = st.text_input("Data size / volume *", value=d.get("data_size",""), placeholder="e.g. 45,000 annotated frames across 12 stores")
        d["date_range"]    = st.text_input("Date range of data", value=d.get("date_range",""), placeholder="e.g. Jan 2024 – Mar 2026")
    with col2:
        d["preprocessing"] = st.text_area("Preprocessing / cleaning steps *", value=d.get("preprocessing",""), height=100,
            placeholder="Frame extraction, deduplication, filtering, augmentation, anonymisation.")
        d["known_biases"]  = st.text_area("Known data biases or gaps *", value=d.get("known_biases",""), height=100,
            placeholder="Class imbalance, geographic skew, seen/unseen store distribution, lighting conditions.")
        d["demographic_notes"] = st.text_area("Demographic representation notes", value=d.get("demographic_notes",""), height=60,
            placeholder="Store formats, SCO vs staffed lanes, geography, retailer mix.")
    section_end()

    # ── 3 · Data Governance (Art. 10) ─────────────────────────────────────────
    if is_high:
        section_header("3 · Data Governance", "Art. 10(2)–10(5)", "required")
        d["data_governance"] = st.text_area(
            "Data governance framework *",
            value=d.get("data_governance",""), height=120,
            placeholder="Describe data management practices: labelling standards, version control, access controls, retention policy, GDPR compliance, consent mechanisms if applicable."
        )
        col1, col2 = st.columns(2)
        with col1:
            d["annotation_methodology"] = st.text_area("Annotation methodology", value=d.get("annotation_methodology",""), height=80,
                placeholder="Annotation guidelines, quality control, inter-annotator agreement, confirmability rules.")
        with col2:
            d["data_version_control"]   = st.text_area("Dataset version control", value=d.get("data_version_control",""), height=80,
                placeholder="Dataset versioning approach, changelog, experiment tracking linkage.")
        section_end()
    else:
        d["data_governance"] = ""

    # ── 4 · Performance Metrics (Art. 15) ─────────────────────────────────────
    section_header("4 · Performance Metrics", "Art. 15 / Annex IV §4", "required" if is_high else "high")
    col1, col2 = st.columns(2)
    with col1:
        d["eval_datasets"]  = st.text_area("Evaluation datasets *", value=d.get("eval_datasets",""), height=100,
            placeholder="Name each split or benchmark. Include seen vs unseen store breakdown.")
        d["metrics_used"]   = st.text_input("Metrics used *", value=d.get("metrics_used",""), placeholder="e.g. Precision, Recall, F1, True Positive Rate, False Positive Rate")
        d["results_summary"] = st.text_area("Results summary *", value=d.get("results_summary",""), height=100,
            placeholder="Key numbers per metric per split. e.g. Seen store recall: 81%, Unseen: 62%.")
    with col2:
        d["disaggregated_eval"]    = st.text_area("Disaggregated evaluation (subgroups)", value=d.get("disaggregated_eval",""), height=100,
            placeholder="Performance by store format, SCO type, lighting condition, product category.")
        d["benchmark_comparisons"] = st.text_area("Benchmark / competitor comparisons", value=d.get("benchmark_comparisons",""), height=60,
            placeholder="Comparison to baseline or competitor (e.g. Everseen bake-off results).")
    section_end()

    # ── 5 · Accuracy, Robustness & Cybersecurity (Art. 15) ───────────────────
    if is_high:
        section_header("5 · Accuracy, Robustness & Cybersecurity", "Art. 15", "required")
        col1, col2 = st.columns(2)
        with col1:
            d["accuracy_robustness"] = st.text_area(
                "Accuracy & robustness measures *",
                value=d.get("accuracy_robustness",""), height=120,
                placeholder="How is accuracy maintained across deployment contexts? Domain shift testing, out-of-distribution evaluation, adversarial testing, stress testing approach."
            )
        with col2:
            d["cybersecurity_measures"] = st.text_area(
                "Cybersecurity measures",
                value=d.get("cybersecurity_measures",""), height=120,
                placeholder="Model access controls, inference endpoint security, adversarial input defences, data poisoning mitigations."
            )
        section_end()
    else:
        d["accuracy_robustness"] = ""

    # ── 6 · Known Limitations (Art. 13) ──────────────────────────────────────
    section_header("6 · Known Limitations", "Art. 13(3)(b)", "high")
    col1, col2 = st.columns(2)
    with col1:
        d["technical_limitations"] = st.text_area("Technical limitations *", value=d.get("technical_limitations",""), height=100,
            placeholder="Input constraints, inference latency, hardware requirements, model size.")
        d["failure_modes"]         = st.text_area("Failure modes", value=d.get("failure_modes",""), height=80,
            placeholder="Known edge cases: occlusion, low light, fast motion, novel product types.")
    with col2:
        d["fairness_concerns"]    = st.text_area("Fairness / bias concerns *", value=d.get("fairness_concerns",""), height=100,
            placeholder="Disparate performance across store formats, demographics, or product categories.")
        d["conditions_degrade"]   = st.text_area("Conditions where performance degrades", value=d.get("conditions_degrade",""), height=80,
            placeholder="Domain shift, adversarial inputs, unseen SCO types, novel store layouts.")
    section_end()

    # ── 7 · Intended Deployment Context (Art. 13) ────────────────────────────
    section_header("7 · Intended Deployment Context", "Art. 13(3)(a)", "high")
    col1, col2 = st.columns(2)
    with col1:
        d["primary_use_cases"]       = st.text_area("Primary use cases *", value=d.get("primary_use_cases",""), height=100,
            placeholder="Specific tasks and scenarios. e.g. Missed scan detection at self-checkout in grocery retail.")
        d["intended_users"]          = st.text_area("Intended users / audience", value=d.get("intended_users",""), height=80,
            placeholder="Loss prevention teams, store operations managers, retail technology integrators.")
    with col2:
        d["recommended_deployment"]  = st.text_area("Recommended deployment environment", value=d.get("recommended_deployment",""), height=80,
            placeholder="Edge device, cloud, GPU/CPU requirements, camera spec, API vs embedded.")
        d["regulatory_notes"]        = st.text_area("Regulatory / compliance notes", value=d.get("regulatory_notes",""), height=80,
            placeholder="EU AI Act risk tier, GDPR applicability, sector-specific regulations.")
    section_end()

    # ── 8 · Human Oversight (Art. 14) ────────────────────────────────────────
    if is_high:
        section_header("8 · Human Oversight Measures", "Art. 14", "required")
        col1, col2 = st.columns(2)
        with col1:
            d["human_oversight"] = st.text_area("Human oversight requirements *", value=d.get("human_oversight",""), height=120,
                placeholder="When must a human review outputs? What are mandatory review workflows? Who has override authority?")
        with col2:
            d["oversight_mechanisms"] = st.text_area("Oversight mechanisms & controls", value=d.get("oversight_mechanisms",""), height=120,
                placeholder="Dashboard monitoring, alert thresholds, human-in-the-loop triggers, escalation paths.")
        section_end()
    else:
        section_header("8 · Human Oversight Measures", "Art. 14", "high")
        d["human_oversight"] = st.text_area("Human oversight requirements *", value=d.get("human_oversight",""), height=100,
            placeholder="When should a human review model outputs? Any mandatory review workflows?")
        section_end()

    # ── 9 · Post-Market Monitoring (Art. 9 / 72) ─────────────────────────────
    if is_high:
        section_header("9 · Post-Market Monitoring", "Art. 9(7) / Art. 72", "required")
        col1, col2 = st.columns(2)
        with col1:
            d["post_market_monitoring"] = st.text_area(
                "Monitoring plan *",
                value=d.get("post_market_monitoring",""), height=120,
                placeholder="How will model performance be tracked in production? KPIs, dashboards, review cadence, drift detection."
            )
        with col2:
            d["monitoring_kpis"]        = st.text_area("Monitoring KPIs & thresholds", value=d.get("monitoring_kpis",""), height=80,
                placeholder="e.g. If precision drops below 70%, trigger review. Weekly recall reporting.")
            d["model_update_policy"]    = st.text_area("Model update / retraining policy", value=d.get("model_update_policy",""), height=80,
                placeholder="When is retraining triggered? Version control, re-validation requirements.")
        section_end()
    else:
        d["post_market_monitoring"] = ""

    # ── 10 · Incident Reporting (Art. 62) ────────────────────────────────────
    if is_high:
        section_header("10 · Incident Reporting", "Art. 62", "required")
        col1, col2 = st.columns(2)
        with col1:
            d["incident_reporting"] = st.text_area(
                "Incident reporting obligations *",
                value=d.get("incident_reporting",""), height=120,
                placeholder="Define 'serious incident'. Reporting timeline (72 hours to market surveillance authority). Who is responsible? Which authority?"
            )
        with col2:
            d["incident_contact"]   = st.text_area("Incident reporting contact", value=d.get("incident_contact",""), height=80,
                placeholder="Internal contact and regulatory authority contact details.")
            d["incident_log"]       = st.text_area("Incident logging approach", value=d.get("incident_log",""), height=80,
                placeholder="How are near-misses and incidents logged? Audit trail approach.")
        section_end()
    else:
        d["incident_reporting"] = ""

    # ── 11 · Conformity Assessment (Art. 43) ─────────────────────────────────
    if is_high:
        section_header("11 · Conformity Assessment", "Art. 43 / Annex IV", "required")
        col1, col2 = st.columns(2)
        with col1:
            d["conformity_assessment_type"] = st.selectbox(
                "Conformity assessment type *",
                ["Internal assessment (Annex VI)", "Third-party assessment", "Notified body assessment", "Not yet conducted", "Not applicable"],
                index=["Internal assessment (Annex VI)", "Third-party assessment", "Notified body assessment", "Not yet conducted", "Not applicable"].index(d.get("conformity_assessment_type", "Not yet conducted"))
            )
            d["notified_body"] = st.text_input("Notified body (if applicable)", value=d.get("notified_body",""), placeholder="Name and ID of notified body")
        with col2:
            d["declaration_of_conformity"] = st.text_area("Declaration of conformity reference", value=d.get("declaration_of_conformity",""), height=80,
                placeholder="Document reference, date, signatory.")
            d["ce_marking_status"] = st.selectbox(
                "CE marking status",
                ["Not applicable", "CE marked", "CE marking in progress", "CE marking not yet sought"],
                index=["Not applicable", "CE marked", "CE marking in progress", "CE marking not yet sought"].index(d.get("ce_marking_status","Not applicable"))
            )
        section_end()
    else:
        d["conformity_assessment_type"] = "Not applicable"

    # ── 12 · Out-of-Scope Uses ───────────────────────────────────────────────
    section_header("12 · Out-of-Scope & Prohibited Uses", "Art. 13(3)(b)(iii)", "high")
    col1, col2 = st.columns(2)
    with col1:
        d["out_of_scope_uses"] = st.text_area("Explicitly out-of-scope uses *", value=d.get("out_of_scope_uses",""), height=100,
            placeholder="Use cases the model was NOT designed for and may perform poorly on.")
        d["misuse_scenarios"]  = st.text_area("Potential misuse scenarios", value=d.get("misuse_scenarios",""), height=80,
            placeholder="Ways the model could be deliberately misused.")
    with col2:
        d["prohibited_uses"]   = st.text_area("Prohibited use cases", value=d.get("prohibited_uses",""), height=100,
            placeholder="Uses explicitly forbidden by licence, ethics policy, or EU AI Act Article 5.")
    section_end()

    # ── 13 · Additional Notes ─────────────────────────────────────────────────
    section_header("13 · Additional Notes", badge="optional")
    d["additional_notes"] = st.text_area("Additional notes (optional)", value=d.get("additional_notes",""), height=80,
        placeholder="Acknowledgements, related papers, funding sources, version history, etc.")
    section_end()

    # ── Compliance Score ──────────────────────────────────────────────────────
    score = compute_score(d, tier)
    score_colour = "#22c55e" if score >= 80 else "#f97316" if score >= 50 else "#dc2626"
    st.markdown(f"""
    <div style="background:white;border:1px solid #e2e8f0;border-radius:10px;padding:1rem 1.5rem;margin:1.5rem 0;">
      <strong>Completeness score: <span style="color:{score_colour}">{score}%</span></strong>
      <div class="score-bar-outer"><div class="score-bar-inner" style="width:{score}%;background:linear-gradient(90deg,{score_colour},{score_colour}99);"></div></div>
      <small style="color:#6b7280;">Fill all required fields (*) to maximise compliance coverage before generating.</small>
    </div>
    """, unsafe_allow_html=True)

    # ── Generate button ───────────────────────────────────────────────────────
    col1, col2 = st.columns([1, 3])
    with col1:
        if st.button("← Back to Risk Wizard"):
            st.session_state.step = "wizard"
            st.rerun()
    with col2:
        generate = st.button("✨ Generate Model Card", type="primary", use_container_width=True)

    if generate:
        api_key = os.environ.get("ANTHROPIC_API_KEY", "")
        if not api_key:
            st.error("ANTHROPIC_API_KEY not found. Set it before launching the app.")
        else:
            st.session_state.data = d
            with st.spinner("Generating your EU AI Act compliant model card…"):
                client = anthropic.Anthropic(api_key=api_key)

                prompt = f"""You are an expert in EU AI Act compliance and ML governance. Generate a comprehensive, professional model card based on the data below.

RISK TIER: {d.get('risk_tier','High')}

INPUT DATA:
{json.dumps(d, indent=2)}

OUTPUT REQUIREMENTS:
- Structure the card with clear sections matching the EU AI Act Article obligations
- For each High Risk section, cite the relevant Article
- Write professionally but accessibly
- Flag any critical gaps (fields left blank that are required for compliance)
- Include a "Compliance Summary" section at the end that lists:
  * Which Article obligations are satisfied
  * Which are partially satisfied
  * Which are missing and the risk of non-compliance
- Format in clean Markdown
- Do NOT invent or fabricate data — only use what is provided. If a field is empty, note it as "Not provided — required for {d.get('risk_tier','')} risk compliance" where applicable.

Generate the full model card now:"""

                response = client.messages.create(
                    model="claude-sonnet-4-20250514",
                    max_tokens=4000,
                    messages=[{"role": "user", "content": prompt}]
                )
                st.session_state.generated = response.content[0].text
                st.session_state.step = "result"
                st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# STEP 3 — RESULT
# ══════════════════════════════════════════════════════════════════════════════
elif st.session_state.step == "result":
    d = st.session_state.data
    tier = d.get("risk_tier", "High")
    colour = risk_colour(tier)
    score = compute_score(d, tier)
    score_colour = "#22c55e" if score >= 80 else "#f97316" if score >= 50 else "#dc2626"

    col1, col2, col3 = st.columns([1, 1, 2])
    with col1:
        if st.button("← Edit Form"):
            st.session_state.step = "form"
            st.rerun()
    with col2:
        if st.button("🔄 Regenerate"):
            st.session_state.step = "form"
            st.rerun()

    st.markdown(f"""
    <div style="background:{colour}15;border:2px solid {colour};border-radius:10px;padding:0.75rem 1.25rem;margin:1rem 0;display:flex;align-items:center;justify-content:space-between;">
      <strong style="color:{colour};">Risk Tier: {tier}</strong>
      <span style="color:{score_colour};font-weight:600;">Completeness: {score}%</span>
    </div>
    """, unsafe_allow_html=True)

    # ── Display generated card ─────────────────────────────────────────────────
    with st.container():
        st.markdown('<div class="generated-card">', unsafe_allow_html=True)
        st.markdown(st.session_state.generated)
        st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.subheader("Export")

    col1, col2, col3 = st.columns(3)

    # ── Markdown download ──────────────────────────────────────────────────────
    with col1:
        md_content = f"# Model Card — {d.get('model_name','Unnamed Model')}\n\n"
        md_content += f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}  \n"
        md_content += f"**Risk Tier:** {tier}  \n"
        md_content += f"**Completeness Score:** {score}%\n\n---\n\n"
        md_content += st.session_state.generated
        st.download_button(
            "⬇️ Download Markdown",
            data=md_content.encode("utf-8"),
            file_name=f"model_card_{d.get('model_name','model').replace(' ','_')}.md",
            mime="text/markdown",
            use_container_width=True,
        )

    # ── HTML download ──────────────────────────────────────────────────────────
    with col2:
        try:
            import markdown as md_lib
            html_body = md_lib.markdown(st.session_state.generated, extensions=["tables", "fenced_code"])
        except ImportError:
            html_body = f"<pre>{st.session_state.generated}</pre>"

        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Model Card — {d.get('model_name','Model')}</title>
<style>
  body {{ font-family: 'Segoe UI', Arial, sans-serif; max-width: 900px; margin: 40px auto; padding: 0 2rem; color: #1a202c; line-height: 1.7; }}
  h1 {{ color: #0f172a; border-bottom: 3px solid #0f4c81; padding-bottom: 0.5rem; }}
  h2 {{ color: #1e3a5f; margin-top: 2rem; }}
  h3 {{ color: #374151; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1rem 0; }}
  th, td {{ border: 1px solid #e2e8f0; padding: 0.6rem 1rem; text-align: left; }}
  th {{ background: #0f172a; color: white; }}
  tr:nth-child(even) {{ background: #f8f9fb; }}
  code {{ background: #f1f5f9; padding: 0.2rem 0.4rem; border-radius: 4px; font-size: 0.9em; }}
  pre {{ background: #f1f5f9; padding: 1rem; border-radius: 8px; overflow-x: auto; }}
  .meta {{ background: #eff6ff; border: 1px solid #bfdbfe; border-radius: 8px; padding: 1rem; margin-bottom: 2rem; font-size: 0.9rem; }}
  .footer {{ margin-top: 3rem; padding-top: 1rem; border-top: 1px solid #e2e8f0; font-size: 0.8rem; color: #9ca3af; }}
</style>
</head>
<body>
<div class="meta">
  <strong>Generated:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M')} &nbsp;|&nbsp;
  <strong>Risk Tier:</strong> {tier} &nbsp;|&nbsp;
  <strong>Completeness Score:</strong> {score}%
</div>
{html_body}
<div class="footer">Generated by CA Project Management Services Ltd — EU AI Act Model Card Generator</div>
</body>
</html>"""
        st.download_button(
            "⬇️ Download HTML",
            data=html_content.encode("utf-8"),
            file_name=f"model_card_{d.get('model_name','model').replace(' ','_')}.html",
            mime="text/html",
            use_container_width=True,
        )

    # ── Word (.docx) download ──────────────────────────────────────────────────
    with col3:
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = DocxDocument()

            # Styles
            style = doc.styles["Normal"]
            style.font.name = "Arial"
            style.font.size = Pt(11)

            # Title
            title_para = doc.add_heading(f"Model Card — {d.get('model_name','Unnamed Model')}", 0)
            title_para.runs[0].font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

            # Meta block
            meta = doc.add_paragraph()
            meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}   |   Risk Tier: {tier}   |   Completeness: {score}%").italic = True

            doc.add_paragraph()

            # Write generated content line by line
            for line in st.session_state.generated.split("\n"):
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("- ") or line.startswith("* "):
                    p = doc.add_paragraph(line[2:], style="List Bullet")
                elif line.startswith("**") and line.endswith("**"):
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip("**"))
                    run.bold = True
                else:
                    # Handle inline bold
                    p = doc.add_paragraph()
                    parts = line.split("**")
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        if i % 2 == 1:
                            run.bold = True

            # Footer
            doc.add_paragraph()
            footer_para = doc.add_paragraph("Generated by CA Project Management Services Ltd — EU AI Act Model Card Generator")
            footer_para.runs[0].font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)
            footer_para.runs[0].font.size = Pt(9)

            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)

            st.download_button(
                "⬇️ Download Word (.docx)",
                data=buf.getvalue(),
                file_name=f"model_card_{d.get('model_name','model').replace(' ','_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except ImportError:
            st.info("Install python-docx for Word export: pip install python-docx")

    # ── Footer ─────────────────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown(
        "<div style='text-align:center;color:#9ca3af;font-size:0.8rem;'>CA Project Management Services Ltd &nbsp;·&nbsp; EU AI Act Model Card Generator &nbsp;·&nbsp; Powered by Claude</div>",
        unsafe_allow_html=True,
    )
